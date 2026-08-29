import random
import logging
import os
import shutil
import io
from typing import List, Optional
from datetime import date, datetime
from pydantic import BaseModel
from fastapi import FastAPI, Depends, HTTPException, BackgroundTasks, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import or_  
from database import get_db
import models
from email_service import send_workflow_email
import cloudinary
import cloudinary.uploader
from cloudinary.utils import cloudinary_url

# 🎯 PURE PYTHON ENGINES (Zero System Dependencies for Render.com)
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# 1. System Logging Configurations
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("AarviProcure")
app = FastAPI(title="Aarvi Encon - Workflow ERP Engine", version="3.1.0")

# 🎯 Configure Cloudinary securely
cloudinary.config(
    cloud_name=os.getenv("CLOUDINARY_CLOUD_NAME"),
    api_key=os.getenv("CLOUDINARY_API_KEY"),
    api_secret=os.getenv("CLOUDINARY_API_SECRET"),
    secure=True
)

UPLOAD_DIR = "storage/quotation_files"
os.makedirs(UPLOAD_DIR, exist_ok=True)
# 🎯 Directory for saved PO HTML templates
os.makedirs("storage/po_templates", exist_ok=True)
app.mount("/storage", StaticFiles(directory="storage"), name="storage")

# 2. Complete CORS Rules
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "https://aarvi-procure-system.vercel.app", "https://procure.aarviencon.com"], 
    allow_credentials=True,
    allow_methods=["*"], 
    allow_headers=["*"], 
)

# -------------------------------------------------------------------
# PYDANTIC INCOMING DATA VALIDATORS
# -------------------------------------------------------------------
class RequisitionRowItem(BaseModel):
    product_description: str
    make_brand: Optional[str] = None
    quantity: int
    purpose: str
    item_type: Optional[str] = "Consumable"  

class CreateRequisitionPayload(BaseModel):
    project_code: str
    project_name: str
    coordinator_id: int
    category: str 
    assigned_site_manager_id: Optional[int] = None
    assigned_project_manager_id: int
    items: List[RequisitionRowItem]
    is_manager_direct_route: Optional[bool] = False 

class DirectPOItemRow(BaseModel):
    product_description: str
    make_brand: Optional[str] = None
    quantity: int
    purpose: str
    item_type: Optional[str] = "Consumable"
    vendor_name: str
    base_total_value: float
    gst_percentage: Optional[float] = 18.0
    net_amount_payable: float
    vendor_address: Optional[str] = ""
    vendor_contact: Optional[str] = ""
    vendor_email: Optional[str] = ""
    time_of_delivery: Optional[str] = "7 Days"
    delivery_address: Optional[str] = ""
    site_contact_person: Optional[str] = ""
    site_contact_phone: Optional[str] = ""
    special_terms: Optional[str] = ""
    quality_remarks: Optional[str] = ""
    file_url: Optional[str] = ""

class DirectPOPayload(BaseModel):
    project_code: str
    project_name: str
    creator_id: int
    creator_name: str
    category: str
    items: List[DirectPOItemRow]    

class UpdateRequisitionItem(BaseModel):
    item_index: int
    product_description: str
    make_brand: Optional[str] = None
    quantity: int
    purpose: str
    is_reimbursable: Optional[bool] = False
    item_type: Optional[str] = "Consumable"  

class ProposeEditsPayload(BaseModel):
    user_name: str
    user_role: str
    remarks: str
    items: List[UpdateRequisitionItem]

class DualApprovalPayload(BaseModel):
    user_name: str
    user_role: str
    items: Optional[List[UpdateRequisitionItem]] = None 

class QuotationRowItem(BaseModel):
    item_index: int
    vendor_name: str
    total_amount: float
    product_description: Optional[str] = None
    make_brand: Optional[str] = None
    quantity: Optional[int] = None
    unit_price: Optional[float] = None
    gst_percentage: Optional[float] = None
    freight_charges: Optional[float] = None
    time_of_delivery: Optional[str] = None
    payment_terms: Optional[str] = None
    contract_start_date: Optional[date] = None
    contract_end_date: Optional[date] = None
    file_url: Optional[str] = None
    special_terms: Optional[str] = None
    quality_remarks: Optional[str] = None
    vendor_address: Optional[str] = None
    vendor_contact: Optional[str] = None
    vendor_email: Optional[str] = None
    delivery_address: Optional[str] = None
    site_contact_person: Optional[str] = None
    site_contact_phone: Optional[str] = None
    base_total_value: Optional[float] = 0.0
    net_amount_payable: Optional[float] = 0.0

class SubmitQuotationsPayload(BaseModel):
    quotations: List[QuotationRowItem]
    items: Optional[List[UpdateRequisitionItem]] = None 

class FinanceApprovalPayload(BaseModel):
    user_name: str
    action: str 
    remarks: Optional[str] = None
    selected_bids: Optional[dict] = None  
    items: Optional[List[UpdateRequisitionItem]] = None 

# 🎯 Payload for saving custom edited PO templates
class SaveTemplatePayload(BaseModel):
    html_content: str


# -------------------------------------------------------------------
# STAGE 0: LIVE PERSONNEL ROUTING
# -------------------------------------------------------------------
@app.get("/api/users/by-role", response_model=List[dict])
def get_active_users_by_role(role: str, db: Session = Depends(get_db)):
    users = db.query(models.User).filter(
        models.User.role == role,
        models.User.is_active == True
    ).order_by(models.User.name.asc()).all()
    return [{"id": u.id, "name": u.name, "email": u.email, "empcode": u.empcode} for u in users]

# -------------------------------------------------------------------
# STAGE 1: SITE COORDINATOR ENTRY GATEWAY
# -------------------------------------------------------------------
@app.post("/api/requisitions", status_code=201)
def raise_material_requisition(
    payload: CreateRequisitionPayload, 
    background_tasks: BackgroundTasks, 
    db: Session = Depends(get_db)
):
    ticket_number = f"REQ-2026-{random.randint(100000, 999999)}"
    
    if payload.is_manager_direct_route:
        initial_status = "Pending Sourcing"
    else:
        initial_status = "Vetting Active" if payload.assigned_site_manager_id else "Pending PM Vetting"
    
    master_ticket = models.MaterialTicket(
        ticket_number=ticket_number,
        project_code=payload.project_code,
        project_name=payload.project_name,
        coordinator_id=payload.coordinator_id,
        category=payload.category,
        assigned_site_manager_id=payload.assigned_site_manager_id,
        assigned_project_manager_id=payload.assigned_project_manager_id,
        status=initial_status
    )
    db.add(master_ticket)
    
    for idx, row in enumerate(payload.items, start=1):
        db_item = models.TicketItem(
            ticket_number=ticket_number,
            item_index=idx,
            product_description=row.product_description,
            make_brand=row.make_brand,
            quantity=row.quantity,
            purpose=row.purpose,
            item_type=row.item_type
        )
        db.add(db_item)
        
    if payload.is_manager_direct_route:
        remarks_text = "Manager Direct Sourcing Request. Routed directly to Purchasing Desk."
    else:
        remarks_text = f"Material Sheet uploaded. Routed to {'Site Manager' if payload.assigned_site_manager_id else 'Project Manager'}."
        
    history = models.TicketHistory(
        ticket_number=ticket_number,
        user_name=f"User ID: {payload.coordinator_id}",
        action_taken="Ticket Raised",
        remarks=remarks_text
    )
    db.add(history)
    db.commit()
    
    target_user_id = (
        payload.assigned_site_manager_id or payload.assigned_project_manager_id
    ) if not payload.is_manager_direct_route else None
    
    if target_user_id:
        target_user = db.query(models.User).filter(models.User.id == target_user_id).first()
        if target_user and target_user.email:
            background_tasks.add_task(
                send_workflow_email,
                recipient_email=target_user.email,
                recipient_name=target_user.name,
                subject="Action Required: New Material Requisition Pending Technical Vetting",
                ticket_number=ticket_number,
                project_name=payload.project_name,
                status=initial_status
            )
    
    return {"ticket_number": ticket_number, "status": initial_status}

@app.post("/api/requisitions/direct-fast-track", status_code=201)
def raise_direct_manager_purchase_order(
    payload: DirectPOPayload, 
    background_tasks: BackgroundTasks, 
    db: Session = Depends(get_db)
):
    ticket_number = f"REQ-2026-{random.randint(100000, 999999)}"
    po_number = f"PO-2026-{random.randint(100000, 999999)}"
    
    master_ticket = models.MaterialTicket(
        ticket_number=ticket_number,
        project_code=payload.project_code,
        project_name=payload.project_name,
        coordinator_id=payload.creator_id,  
        category=payload.category,
        assigned_project_manager_id=payload.creator_id,
        status="Awaiting Digital Signature"
    )
    db.add(master_ticket)
    
    for idx, row in enumerate(payload.items, start=1):
        db_item = models.TicketItem(
            ticket_number=ticket_number,
            item_index=idx,
            product_description=row.product_description,
            make_brand=row.make_brand,
            quantity=row.quantity,
            purpose=row.purpose,
            item_type=row.item_type
        )
        db.add(db_item)
        
        db_quote = models.Quotation(
            ticket_number=ticket_number,
            item_index=idx,
            vendor_name=row.vendor_name,
            total_amount=row.net_amount_payable,
            product_description=row.product_description,
            make_brand=row.make_brand,
            quantity=row.quantity,
            unit_price=round(row.base_total_value / row.quantity, 2) if row.quantity else row.base_total_value,
            gst_percentage=row.gst_percentage,
            base_total_value=row.base_total_value,
            net_amount_payable=row.net_amount_payable,
            time_of_delivery=row.time_of_delivery,
            vendor_address=row.vendor_address,
            vendor_contact=row.vendor_contact,
            vendor_email=row.vendor_email,
            delivery_address=row.delivery_address,
            site_contact_person=row.site_contact_person,
            site_contact_phone=row.site_contact_phone,
            special_terms=row.special_terms,
            quality_remarks=row.quality_remarks,
            file_url=row.file_url,
            is_selected=True  
        )
        db.add(db_quote)
        
    new_po = models.PurchaseOrder(
        po_number=po_number,
        ticket_number=ticket_number,
        pdf_url=f"/storage/aarvi_pos/{po_number}.pdf"
    )
    db.add(new_po)
    
    db.add(models.TicketHistory(
        ticket_number=ticket_number,
        user_name=payload.creator_name,
        action_taken="Direct Fast-Track Execution",
        remarks=f"Manager Direct Request. Auto-compiled technical biddings and issued Draft PO {po_number} without intermediary clearance steps."
    ))
    
    db.commit()
    
    creator_user = db.query(models.User).filter(models.User.id == payload.creator_id).first()
    if creator_user and creator_user.email:
        background_tasks.add_task(
            send_workflow_email,
            recipient_email=creator_user.email,
            recipient_name=creator_user.name,
            subject=f"Fast-Track PO {po_number} Issued & Pending Digital Seal",
            ticket_number=ticket_number,
            project_name=payload.project_name,
            status="Awaiting Digital Signature"
        )
        
    return {"ticket_number": ticket_number, "po_number": po_number, "status": "Awaiting Digital Signature"}    

# -------------------------------------------------------------------
# STAGE 2: DUAL-SIGNATURE NEGOTIATION LOOP 
# -------------------------------------------------------------------
@app.put("/api/requisitions/{ticket_number}/propose-edits")
def propose_ticket_edits(
    ticket_number: str, 
    payload: ProposeEditsPayload, 
    background_tasks: BackgroundTasks, 
    db: Session = Depends(get_db)
):
    ticket = db.query(models.MaterialTicket).filter(models.MaterialTicket.ticket_number == ticket_number).first()
    if not ticket: raise HTTPException(status_code=404, detail="Requisition not found.")
        
    db.query(models.TicketItem).filter(models.TicketItem.ticket_number == ticket_number).delete()
    for row in payload.items:
        db.add(models.TicketItem(
            ticket_number=ticket_number,
            item_index=row.item_index,
            product_description=row.product_description,
            make_brand=row.make_brand,
            quantity=row.quantity,
            purpose=row.purpose,
            is_reimbursable=row.is_reimbursable,
            item_type=row.item_type  
        ))
        
    if payload.user_role in ["Site Manager", "Project Manager"]:
        ticket.status = "Awaiting Coordinator Sign-Off"
    else:
        ticket.status = "Vetting Active" if ticket.assigned_site_manager_id else "Pending PM Vetting"
    
    db.add(models.TicketHistory(
        ticket_number=ticket_number,
        user_name=payload.user_name,
        action_taken="Proposed Counter-Edits",
        remarks=payload.remarks
    ))
    db.commit()
    
    coordinator = db.query(models.User).filter(models.User.id == ticket.coordinator_id).first()
    if coordinator and coordinator.email:
        background_tasks.add_task(
            send_workflow_email,
            recipient_email=coordinator.email,
            recipient_name=coordinator.name,
            subject=f"Action Required: Query / Counter-Edits Flagged on {ticket_number}",
            ticket_number=ticket_number,
            project_name=ticket.project_name,
            status=ticket.status
        )
        
    return {"ticket_number": ticket_number, "status": ticket.status}

@app.put("/api/requisitions/{ticket_number}/approve")
def dual_sign_approve(
    ticket_number: str, 
    payload: DualApprovalPayload, 
    background_tasks: BackgroundTasks, 
    db: Session = Depends(get_db)
):
    ticket = db.query(models.MaterialTicket).filter(models.MaterialTicket.ticket_number == ticket_number).first()
    if not ticket: raise HTTPException(status_code=404, detail="Requisition not found.")
    
    if payload.items:
        db.query(models.TicketItem).filter(models.TicketItem.ticket_number == ticket_number).delete()
        for row in payload.items:
            db.add(models.TicketItem(
                ticket_number=ticket_number,
                item_index=row.item_index,
                product_description=row.product_description,
                make_brand=row.make_brand,
                quantity=row.quantity,
                purpose=row.purpose,
                is_reimbursable=row.is_reimbursable,
                item_type=row.item_type
            ))
            
    remarks_text = f"List Approved & Locked by {payload.user_role}."
    
    if payload.user_role in ["Site Manager", "Project Manager"]:
        ticket.status = "Pending Sourcing"
        remarks_text += " Technical Vetting complete. Dispatched directly to Purchasing Desk."
            
    elif payload.user_role == "Site Coordinator":
        ticket.status = "Pending Sourcing"
        remarks_text += " Coordinator accepted Manager's counter-edits. Dispatched directly to Purchasing Desk."
            
    db.add(models.TicketHistory(
        ticket_number=ticket_number,
        user_name=payload.user_name,
        action_taken="Explicit Sign-Off Applied",
        remarks=remarks_text
    ))
    db.commit()
    
    if ticket.status == "Pending Sourcing":
        purchase_execs = db.query(models.User).filter(models.User.role == "Purchase Executive", models.User.is_active == True).all()
        for pe in purchase_execs:
            if pe.email:
                background_tasks.add_task(
                    send_workflow_email,
                    recipient_email=pe.email,
                    recipient_name=pe.name,
                    subject=f"Action Required: Vetting Cleared - Source Bids for {ticket_number}",
                    ticket_number=ticket_number,
                    project_name=ticket.project_name,
                    status="Pending Sourcing"
                )
                
    return {"ticket_number": ticket_number, "status": ticket.status}

# -------------------------------------------------------------------
# STAGE 3: PURCHASING DESK QUOTATION ATTACHMENT
# -------------------------------------------------------------------
@app.post("/api/upload/quotation")
async def upload_quotation_document(
    ticket_number: str,
    item_index: int,
    option_index: int,
    file: UploadFile = File(...)
):
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in [".pdf", ".doc", ".docx"]:
        raise HTTPException(status_code=400, detail="Unsupported format. Only PDF, DOC, and DOCX are allowed.")
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    systematic_name = f"QUOTE_{ticket_number}_ROW{item_index}_OPT{option_index}_{timestamp}{ext}"
    target_destination = os.path.join(UPLOAD_DIR, systematic_name)
    
    try:
        with open(target_destination, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Disk write failure: {str(e)}")
        
    return {"file_url": f"/storage/quotation_files/{systematic_name}"}

@app.post("/api/requisitions/{ticket_number}/quotations")
def attach_vendor_quotations(
    ticket_number: str, 
    payload: SubmitQuotationsPayload, 
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db)
):
    ticket = db.query(models.MaterialTicket).filter(models.MaterialTicket.ticket_number == ticket_number).first()
    if not ticket: raise HTTPException(status_code=404, detail="Active requisition sheet not found.")
        
    if payload.items:
        for row in payload.items:
            db.query(models.TicketItem).filter(
                models.TicketItem.ticket_number == ticket_number,
                models.TicketItem.item_index == row.item_index
            ).update({
                "item_type": row.item_type,
                "quantity": row.quantity
            })
        
    highest_landed_total = 0.0
    any_unit_price_exceeds_2_5l = False 
    
    for quote in payload.quotations:
        db_quote = models.Quotation(
            ticket_number=ticket_number,
            item_index=quote.item_index,
            vendor_name=quote.vendor_name,
            total_amount=quote.total_amount,
            product_description=quote.product_description,
            make_brand=quote.make_brand,
            quantity=quote.quantity,
            unit_price=quote.unit_price,
            gst_percentage=quote.gst_percentage,
            freight_charges=quote.freight_charges,
            time_of_delivery=quote.time_of_delivery,
            payment_terms=quote.payment_terms,
            contract_start_date=quote.contract_start_date,
            contract_end_date=quote.contract_end_date,
            file_url=getattr(quote, "file_url", "No attachment provided"),
            special_terms=quote.special_terms,
            quality_remarks=quote.quality_remarks,
            vendor_address=quote.vendor_address,
            vendor_contact=quote.vendor_contact,
            vendor_email=quote.vendor_email,
            delivery_address=quote.delivery_address,
            site_contact_person=quote.site_contact_person,
            site_contact_phone=quote.site_contact_phone,
            base_total_value=quote.base_total_value,
            net_amount_payable=quote.net_amount_payable
        )
        db.add(db_quote)
        
        single_unit_price = float(quote.unit_price or 0.0)
        if single_unit_price > 250000:
            any_unit_price_exceeds_2_5l = True

        if quote.vendor_name:
            clean_name = quote.vendor_name.strip()
            existing_vendor = db.query(models.Vendor).filter(models.Vendor.name == clean_name).first()
            if not existing_vendor:
                new_vendor = models.Vendor(
                    name=clean_name,
                    address=quote.vendor_address or "",
                    contact_number=quote.vendor_contact or "",
                    email=quote.vendor_email or "",
                    is_active=True
                )
                db.add(new_vendor)
        
        if quote.total_amount > highest_landed_total:
            highest_landed_total = quote.total_amount
            
    exceeds_10l_total = highest_landed_total > 1000000
    requires_director_review = exceeds_10l_total or any_unit_price_exceeds_2_5l

    if not requires_director_review:
        ticket.status = "Pending Project Manager"
        routing_msg = f"Order matrix value (Highest Total: ₹{highest_landed_total:,.2f}) routed directly to Project Manager for mandatory clearance."
        
        target_pm = db.query(models.User).filter(models.User.id == ticket.assigned_project_manager_id).first()
        if target_pm and target_pm.email:
            background_tasks.add_task(
                send_workflow_email,
                recipient_email=target_pm.email,
                recipient_name=target_pm.name,
                subject=f"Commercial Approval Required: Quotations Attached for {ticket_number}",
                ticket_number=ticket_number,
                project_name=ticket.project_name,
                status=ticket.status
            )
            
    else: 
        ticket.status = "Pending Director"
        reasons = []
        if exceeds_10l_total:
            reasons.append(f"Highest Order Total (₹{highest_landed_total:,.2f}) > ₹10L")
        if any_unit_price_exceeds_2_5l:
            reasons.append("One or more items have an individual Unit Price exceeding ₹2.5L")
            
        routing_msg = f"High-value corporate order routed to Executive Director Board ({', '.join(reasons)})."
        
        directors = db.query(models.User).filter(models.User.role == "Director", models.User.is_active == True).all()
        for director in directors:
            if director.email:
                background_tasks.add_task(
                    send_workflow_email,
                    recipient_email=director.email,
                    recipient_name=director.name,
                    subject=f"High-Value Board Clearance Required for {ticket_number}",
                    ticket_number=ticket_number,
                    project_name=ticket.project_name,
                    status=ticket.status
                )
                
    db.add(models.TicketHistory(
        ticket_number=ticket_number,
        user_name="Procurement Desk Officer",
        action_taken="Quotations Processed",
        remarks=routing_msg
    ))
    db.commit()
    return {"ticket_number": ticket_number, "status": ticket.status}

# -------------------------------------------------------------------
# STAGE 4 & 5: MANAGEMENT SIGN-OFF
# -------------------------------------------------------------------
@app.post("/api/requisitions/{ticket_number}/action")
def process_financial_signoff(
    ticket_number: str, 
    payload: FinanceApprovalPayload, 
    background_tasks: BackgroundTasks, 
    db: Session = Depends(get_db)
):
    ticket = db.query(models.MaterialTicket).filter(models.MaterialTicket.ticket_number == ticket_number).first()
    if not ticket: raise HTTPException(status_code=404, detail="Requisition not found.")
    
    if payload.action == "Approve":
        if not payload.selected_bids:
            raise HTTPException(status_code=400, detail="You must select a winning bid for the items to approve.")
            
        if payload.items:
            for row in payload.items:
                db.query(models.TicketItem).filter(
                    models.TicketItem.ticket_number == ticket_number,
                    models.TicketItem.item_index == row.item_index
                ).update({"is_reimbursable": row.is_reimbursable,
                          "item_type": row.item_type})
                
        for item_idx_str, winning_vendor in payload.selected_bids.items():
            item_idx = int(item_idx_str)
            db.query(models.Quotation).filter(
                models.Quotation.ticket_number == ticket_number,
                models.Quotation.item_index == item_idx,
                models.Quotation.vendor_name == winning_vendor
            ).update({"is_selected": True})
            
        ticket.status = "Awaiting Digital Signature"
        po_number = f"PO-2026-{random.randint(100000, 999999)}"
        
        new_po = models.PurchaseOrder(
            po_number=po_number,
            ticket_number=ticket_number,
            pdf_url=f"/storage/aarvi_pos/{po_number}.pdf"
        )
        db.add(new_po)
        remarks_text = f"Budget cleared by {payload.user_name}. Winning vendor bids locked. Draft PO template {po_number} generated and sent to Purchasing Department."
        
        purchase_execs = db.query(models.User).filter(models.User.role == "Purchase Executive", models.User.is_active == True).all()
        for pe in purchase_execs:
            if pe.email:
                background_tasks.add_task(
                    send_workflow_email,
                    recipient_email=pe.email,
                    recipient_name=pe.name,
                    subject=f"Budget Cleared: Draft PO for {ticket_number} Ready for Digital Signature",
                    ticket_number=ticket_number,
                    project_name=ticket.project_name,
                    status="Awaiting Digital Signature"
                )
                
    elif payload.action == "Raise Query":
        ticket.status = "Query Raised"
        remarks_text = f"Query flagged by {payload.user_name}: {payload.remarks}"
        
        coordinator = db.query(models.User).filter(models.User.id == ticket.coordinator_id).first()
        if coordinator and coordinator.email:
            background_tasks.add_task(
                send_workflow_email,
                recipient_email=coordinator.email,
                recipient_name=coordinator.name,
                subject=f"Action Required: Query Flagged by Management on {ticket_number}",
                ticket_number=ticket_number,
                project_name=ticket.project_name,
                status="Query Raised"
            )
            
    else:
        raise HTTPException(status_code=400, detail="Invalid operational action type.")
        
    db.add(models.TicketHistory(
        ticket_number=ticket_number,
        user_name=payload.user_name,
        action_taken=payload.action,
        remarks=remarks_text
    ))
    db.commit()
    return {"ticket_number": ticket_number, "status": ticket.status}

# -------------------------------------------------------------------
# STAGE 6: PURCHASE ORDER DIGITAL SIGNATURE COMPLETION
# -------------------------------------------------------------------
@app.post("/api/purchase-orders/{po_number}/sign")
def sign_and_finalize_purchase_order(po_number: str, payload: DualApprovalPayload, db: Session = Depends(get_db)):
    po = db.query(models.PurchaseOrder).filter(models.PurchaseOrder.po_number == po_number).first()
    if not po:
        raise HTTPException(status_code=404, detail="Purchase Order records not found.")
        
    ticket = db.query(models.MaterialTicket).filter(models.MaterialTicket.ticket_number == po.ticket_number).first()
    if ticket:
        ticket.status = "Approved"
        
    db.add(models.TicketHistory(
        ticket_number=po.ticket_number,
        user_name=payload.user_name,
        action_taken="Explicit Sign-Off Applied", 
        remarks=f"PO Sealed & Dispatched: Purchase Order {po_number} digitally verified, sealed, and released for logistics tracking by {payload.user_name}."
    ))
    db.commit()
    return {"po_number": po_number, "status": "Approved"}

# -------------------------------------------------------------------
# SYNCHRONIZATION ENDPOINTS
# -------------------------------------------------------------------
@app.get("/api/requisitions/pending-vetting/{manager_id}", response_model=List[dict])
def get_pending_vetting_tickets(manager_id: int, db: Session = Depends(get_db)):
    tickets = db.query(models.MaterialTicket).filter(
        or_(
            (models.MaterialTicket.assigned_site_manager_id == manager_id) & 
            (models.MaterialTicket.status.in_(["Vetting Active", "Approved by Coordinator"])),
            
            (models.MaterialTicket.assigned_project_manager_id == manager_id) & 
            (models.MaterialTicket.status == "Pending PM Vetting"),
            
            (models.MaterialTicket.assigned_site_manager_id == None) & 
            (models.MaterialTicket.status.in_(["Vetting Active", "Approved by Coordinator"]))
        )
    ).order_by(models.MaterialTicket.created_at.desc()).all()
    
    response = []
    for t in tickets:
        coordinator = db.query(models.User).filter(models.User.id == t.coordinator_id).first()
        coordinator_name = coordinator.name if coordinator else "Unknown Coordinator"
        
        response.append({
            "ticket_number": t.ticket_number, 
            "project_code": t.project_code, 
            "project_name": t.project_name, 
            "status": t.status,
            "category": t.category,
            "raised_by": coordinator_name
        })
        
    return response

@app.get("/api/requisitions/pending-handshake/{coordinator_id}", response_model=List[dict])
def get_coordinator_handshake_queue(coordinator_id: int, db: Session = Depends(get_db)):
    tickets = db.query(models.MaterialTicket).filter(
        models.MaterialTicket.coordinator_id == coordinator_id,
        models.MaterialTicket.status.in_(["Awaiting Coordinator Sign-Off", "Approved by Manager"])
    ).order_by(models.MaterialTicket.created_at.desc()).all()
    
    return [{"ticket_number": t.ticket_number, "project_name": t.project_name, "status": t.status, "category": t.category} for t in tickets]

@app.get("/api/requisitions/{ticket_number}/items", response_model=List[dict])
def get_ticket_line_items(ticket_number: str, db: Session = Depends(get_db)):
    items = db.query(models.TicketItem).filter(models.TicketItem.ticket_number == ticket_number).order_by(models.TicketItem.item_index.asc()).all()
    return [{"item_index": i.item_index, "product_description": i.product_description, "make_brand": i.make_brand, "quantity": i.quantity, "purpose": i.purpose, "is_reimbursable": i.is_reimbursable,"item_type": getattr(i, 'item_type', 'Consumable')} for i in items]

@app.get("/api/requisitions/{ticket_number}/history", response_model=List[dict])
def get_ticket_history_logs(ticket_number: str, db: Session = Depends(get_db)):
    logs = db.query(models.TicketHistory).filter(models.TicketHistory.ticket_number == ticket_number).order_by(models.TicketHistory.timestamp.desc()).all()
    return [{"user_name": l.user_name, "action_taken": l.action_taken, "remarks": l.remarks, "timestamp": str(l.timestamp)} for l in logs]

@app.get("/api/requisitions/pending-purchase-approval", response_model=List[dict])
def get_pending_purchase_approval_tickets(db: Session = Depends(get_db)):
    tickets = db.query(models.MaterialTicket).filter(
        models.MaterialTicket.status == "Pending Purchase Approval"
    ).order_by(models.MaterialTicket.created_at.desc()).all()
    
    return [{"ticket_number": t.ticket_number, "project_code": t.project_code, "project_name": t.project_name, "status": t.status, "category": t.category} for t in tickets]

@app.get("/api/requisitions/coordinator-history/{coordinator_id}", response_model=List[dict])
def get_coordinator_completed_history(coordinator_id: int, db: Session = Depends(get_db)):
    tickets = db.query(models.MaterialTicket).filter(
        models.MaterialTicket.coordinator_id == coordinator_id,
        models.MaterialTicket.status.notin_(["Vetting Active", "Awaiting Coordinator Sign-Off"])
    ).order_by(models.MaterialTicket.created_at.desc()).all()
    
    response = []
    for t in tickets:
        log = db.query(models.TicketHistory).filter(
            models.TicketHistory.ticket_number == t.ticket_number,
            models.TicketHistory.action_taken.in_(["Ticket Raised", "Explicit Sign-Off Applied"])
        ).order_by(models.TicketHistory.id.desc()).first()
        
        response.append({
            "ticket_number": t.ticket_number,
            "project_code": t.project_code,
            "project_name": t.project_name,
            "status": t.status,
            "action_date": str(log.timestamp.strftime('%d-%m-%Y %H:%M')) if log else "Date Unavailable"
        })
    return response

@app.get("/api/requisitions/manager-history/{manager_id}", response_model=List[dict])
def get_manager_vetted_history_ledger(manager_id: int, db: Session = Depends(get_db)):
    tickets = db.query(models.MaterialTicket).filter(
        or_(
            models.MaterialTicket.assigned_site_manager_id == manager_id,
            models.MaterialTicket.assigned_site_manager_id == None
        ),
        models.MaterialTicket.status.notin_(["Vetting Active", "Awaiting Coordinator Sign-Off", "Approved by Manager"])
    ).order_by(models.MaterialTicket.created_at.desc()).all()
    
    response = []
    for t in tickets:
        log = db.query(models.TicketHistory).filter(
            models.TicketHistory.ticket_number == t.ticket_number,
            models.TicketHistory.action_taken == "Explicit Sign-Off Applied"
        ).order_by(models.TicketHistory.id.desc()).first()
        
        response.append({
            "ticket_number": t.ticket_number,
            "project_code": t.project_code,
            "project_name": t.project_name,
            "status": t.status,
            "category": t.category,
            "created_at": str(t.created_at.strftime('%d-%m-%Y %H:%M')) if t.created_at else "N/A",
            "action_date": str(log.timestamp.strftime('%d-%m-%Y %H:%M')) if log else "Date Unavailable"
        })
    return response

@app.get("/api/requisitions/pending-sourcing", response_model=List[dict])
def get_pending_sourcing_tickets(db: Session = Depends(get_db)):
    tickets = db.query(models.MaterialTicket).filter(
        models.MaterialTicket.status == "Pending Sourcing"
    ).order_by(models.MaterialTicket.created_at.desc()).all()
    
    return [{"ticket_number": t.ticket_number, "project_code": t.project_code, "project_name": t.project_name, "status": t.status, "category": t.category} for t in tickets]

@app.get("/api/requisitions/purchase-history", response_model=List[dict])
def get_purchase_history(db: Session = Depends(get_db)):
    tickets = db.query(models.MaterialTicket).filter(
        models.MaterialTicket.status.in_([
            "Pending Project Manager", "Pending Director", "Awaiting Digital Signature", 
            "Approved", "Dispatched", "PI Pending PM Approval", "PI Approved - Sent to Accounts",
            "Partially Delivered", "Material Discrepancy Raised", "Delivered - GRN Logged"
        ])
    ).order_by(models.MaterialTicket.created_at.desc()).all()
    
    response = []
    for t in tickets:
        log = db.query(models.TicketHistory).filter(
            models.TicketHistory.ticket_number == t.ticket_number,
            models.TicketHistory.action_taken.in_(["Quotations Processed", "PO Digitally Signed", "Explicit Sign-Off Applied", "Proforma Invoice Uploaded"])
        ).order_by(models.TicketHistory.id.desc()).first()
        
        response.append({
            "ticket_number": t.ticket_number,
            "project_code": t.project_code,
            "project_name": t.project_name,
            "status": t.status,
            "action_date": str(log.timestamp.strftime('%d-%m-%Y %H:%M')) if log else "Date Unavailable",
            "category": t.category
        })
    return response

@app.get("/api/requisitions/pending-management-approval/{manager_id}", response_model=List[dict])
def get_pending_management_approval_tickets(manager_id: int, db: Session = Depends(get_db)):
    user = db.query(models.User).filter(models.User.id == manager_id).first()
    is_director = getattr(user, "role", "") == "Director"
    if is_director:
        tickets = db.query(models.MaterialTicket).filter(
            models.MaterialTicket.status.in_(["Pending Director", "Query Raised"])
        ).order_by(models.MaterialTicket.created_at.desc()).all()
    else:
        tickets = db.query(models.MaterialTicket).filter(
            or_(
                models.MaterialTicket.assigned_project_manager_id == manager_id,
                models.MaterialTicket.assigned_project_manager_id == None
            ),
            models.MaterialTicket.status.in_(["Pending Project Manager", "Query Raised", "PI Pending PM Approval"])
        ).order_by(models.MaterialTicket.created_at.desc()).all()
        
    return [{"ticket_number": t.ticket_number, "project_code": t.project_code, "project_name": t.project_name, "status": t.status} for t in tickets]
   
@app.get("/api/requisitions/pm-history/{manager_id}", response_model=List[dict])
def get_pm_history(manager_id: int, db: Session = Depends(get_db)):
    tickets = db.query(models.MaterialTicket).filter(
        or_(
            models.MaterialTicket.assigned_project_manager_id == manager_id,
            models.MaterialTicket.assigned_project_manager_id == None
        ),
        models.MaterialTicket.status.in_([
            "Pending Director", "Awaiting Digital Signature", "Approved", 
            "Dispatched", "PI Pending PM Approval", "PI Approved - Sent to Accounts",
            "Partially Delivered", "Material Discrepancy Raised", "Delivered - GRN Logged"
        ])
    ).order_by(models.MaterialTicket.created_at.desc()).all()
    
    response = []
    for t in tickets:
        log = db.query(models.TicketHistory).filter(
            models.TicketHistory.ticket_number == t.ticket_number,
            models.TicketHistory.action_taken == "Approve"
        ).order_by(models.TicketHistory.id.desc()).first()
        
        response.append({
            "ticket_number": t.ticket_number,
            "project_code": t.project_code,
            "project_name": t.project_name,
            "status": t.status,
            "approval_date": str(log.timestamp.strftime('%d-%m-%Y %H:%M')) if log else "Date Unavailable"
        })
    return response

@app.get("/api/requisitions/director-history", response_model=List[dict])
def get_director_history(db: Session = Depends(get_db)):
    tickets = db.query(models.MaterialTicket).filter(
        models.MaterialTicket.status.in_(["Awaiting Digital Signature", "Approved", "Dispatched"])
    ).order_by(models.MaterialTicket.created_at.desc()).all()
    
    response = []
    for t in tickets:
        log = db.query(models.TicketHistory).filter(
            models.TicketHistory.ticket_number == t.ticket_number,
            models.TicketHistory.action_taken == "Approve"
        ).order_by(models.TicketHistory.id.desc()).first()
        
        response.append({
            "ticket_number": t.ticket_number,
            "project_code": t.project_code,
            "project_name": t.project_name,
            "status": t.status,
            "approval_date": str(log.timestamp.strftime('%d-%m-%Y %H:%M')) if log else "Date Unavailable"
        })
    return response

@app.get("/api/requisitions/{ticket_number}/quotations", response_model=List[dict])
def get_ticket_vendor_quotations(ticket_number: str, db: Session = Depends(get_db)):
    quotes = db.query(models.Quotation).filter(models.Quotation.ticket_number == ticket_number).all()
    return [
        {
            "item_index": q.item_index,
            "vendor_name": q.vendor_name,
            "total_amount": float(q.total_amount),
            "product_description": q.product_description,
            "make_brand": q.make_brand,
            "quantity": q.quantity,
            "time_of_delivery": q.time_of_delivery,
            "special_terms": q.special_terms,
            "is_selected": q.is_selected,
            "quality_remarks": q.quality_remarks,
            "vendor_address": q.vendor_address,
            "vendor_contact": q.vendor_contact,
            "vendor_email": q.vendor_email,
            "delivery_address": q.delivery_address,
            "site_contact_person": q.site_contact_person,
            "site_contact_phone": q.site_contact_phone,
            "base_total_value": float(q.base_total_value or 0),
            "net_amount_payable": float(q.net_amount_payable or 0)
        } for q in quotes
    ]

@app.get("/api/purchase-orders/pending-signature", response_model=List[dict])
def get_purchase_orders_awaiting_signature(db: Session = Depends(get_db)):
    orders = db.query(
        models.PurchaseOrder, 
        models.MaterialTicket
    ).join(
        models.MaterialTicket, models.PurchaseOrder.ticket_number == models.MaterialTicket.ticket_number
    ).filter(models.MaterialTicket.status == "Awaiting Digital Signature").all()
    
    response = []
    for po_obj, ticket_obj in orders:
        winning_quotes = db.query(models.Quotation).filter(
            models.Quotation.ticket_number == po_obj.ticket_number,
            models.Quotation.is_selected == True
        ).all()
        
        grand_total = sum(q.total_amount for q in winning_quotes)
        vendor_names = list(set(q.vendor_name for q in winning_quotes))
        primary_vendor = vendor_names[0] if vendor_names else "Pending Vendor Linking"
        if len(vendor_names) > 1:
            primary_vendor += f" (+{len(vendor_names)-1} more)"
            
        response.append({
            "po_number": po_obj.po_number,
            "ticket_number": po_obj.ticket_number,
            "vendor_name": primary_vendor,
            "grand_total": float(grand_total),
            "project_name": ticket_obj.project_name,
            "project_code": ticket_obj.project_code,
            "status": "Awaiting Digital Signature",
            "approved_by": "Pending Executive Seal",
            "category": ticket_obj.category 
        })
        
    return response

# -------------------------------------------------------------------
# 💾 SAVE & LOAD EDITED PO HTML TEMPLATES (Supports PO or REQ Ref)
# -------------------------------------------------------------------
@app.put("/api/purchase-orders/{ref_id}/template")
def save_po_template(ref_id: str, payload: SaveTemplatePayload):
    file_path = os.path.join("storage/po_templates", f"{ref_id}.html")
    try:
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(payload.html_content)
        return {"message": "PO Template saved successfully."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save template: {str(e)}")

@app.get("/api/purchase-orders/{ref_id}/template")
def get_po_template(ref_id: str):
    file_path = os.path.join("storage/po_templates", f"{ref_id}.html")
    if os.path.exists(file_path):
        with open(file_path, "r", encoding="utf-8") as f:
            return {"html_content": f.read()}
    return {"html_content": None}

# -------------------------------------------------------------------
# 🎯 1. NATIVE PDF ENDPOINT (Pure Python - Full T&C Template)
# -------------------------------------------------------------------
@app.get("/api/purchase-orders/{po_number}/download-pdf")
def generate_native_vector_pdf(po_number: str, db: Session = Depends(get_db)):
    po = db.query(models.PurchaseOrder).filter(models.PurchaseOrder.po_number == po_number).first()
    if not po: raise HTTPException(status_code=404, detail="Purchase Order not found.")
        
    ticket = db.query(models.MaterialTicket).filter(models.MaterialTicket.ticket_number == po.ticket_number).first()
    winning_quotes = db.query(models.Quotation).filter(models.Quotation.ticket_number == po.ticket_number, models.Quotation.is_selected == True).all()
    if not winning_quotes:
        winning_quotes = db.query(models.Quotation).filter(models.Quotation.ticket_number == po.ticket_number).all()
        
    primary_quote = winning_quotes[0] if winning_quotes else None
    
    # 🎯 FIX: Correctly handle 'None' values with strict fallbacks
    vendor_name = getattr(primary_quote, 'vendor_name', None) or "N/A"
    vendor_address = getattr(primary_quote, 'vendor_address', None) or "Address Not Provided"
    vendor_contact = getattr(primary_quote, 'vendor_contact', None) or "N/A"
    vendor_email = getattr(primary_quote, 'vendor_email', None) or "N/A"
    payment_terms = getattr(primary_quote, 'payment_terms', None) or "100% Payment shall be paid after receipt of material at site."
    time_of_delivery = getattr(primary_quote, 'time_of_delivery', None) or "2-3 Days"
    site_contact = getattr(primary_quote, 'site_contact_person', None) or "Site Coordinator"
    site_phone = getattr(primary_quote, 'site_contact_phone', None) or "N/A"
    delivery_address = getattr(primary_quote, 'delivery_address', None) or "Address Pending"
    contract_date = getattr(primary_quote, 'contract_start_date', None)
    contract_date_str = contract_date.strftime('%d/%m/%Y') if contract_date else "Recently Submitted"
    project_name = ticket.project_name if ticket else "N/A"

    base_total = sum(float(q.base_total_value or 0) for q in winning_quotes)
    net_total = sum(float(q.net_amount_payable or q.total_amount or 0) for q in winning_quotes)
    gst_adj = net_total - base_total

    def number_to_words(num):
        if num == 0: return 'Zero'
        ones = ['', 'One ', 'Two ', 'Three ', 'Four ', 'Five ', 'Six ', 'Seven ', 'Eight ', 'Nine ', 'Ten ', 'Eleven ', 'Twelve ', 'Thirteen ', 'Fourteen ', 'Fifteen ', 'Sixteen ', 'Seventeen ', 'Eighteen ', 'Nineteen ']
        tens = ['', '', 'Twenty ', 'Thirty ', 'Forty ', 'Fifty ', 'Sixty ', 'Seventy ', 'Eighty ', 'Ninety ']
        def convert_less_thousand(n):
            s = ''
            if n >= 100: s += ones[int(n // 100)] + 'Hundred '; n %= 100
            if n >= 20: s += tens[int(n // 10)]; n %= 10
            if n > 0: s += ones[int(n)]
            return s
        str_val = ''
        crore = int(num // 10000000); num %= 10000000
        lakh = int(num // 100000); num %= 100000
        thousand = int(num // 1000); num %= 1000
        if crore > 0: str_val += convert_less_thousand(crore) + 'Crore '
        if lakh > 0: str_val += convert_less_thousand(lakh) + 'Lakh '
        if thousand > 0: str_val += convert_less_thousand(thousand) + 'Thousand '
        if num > 0: str_val += convert_less_thousand(num)
        return str_val.strip() + ' Only'

    amount_in_words = f"Rupees {number_to_words(int(round(net_total)))}"

    table_rows_html = ""
    for idx, q in enumerate(winning_quotes, start=1):
        qty = q.quantity or 1
        unit_rate = (q.base_total_value or 0) / qty
        desc = f"{q.product_description} ({q.make_brand})" if q.make_brand else (q.product_description or 'Item')
        table_rows_html += f"""
        <tr>
            <td style="text-align: center; border: 1px solid #94a3b8; padding: 5px;">0{idx}</td>
            <td style="border: 1px solid #94a3b8; padding: 5px; font-weight: bold;">{desc}</td>
            <td style="text-align: center; border: 1px solid #94a3b8; padding: 5px;">{qty} Nos</td>
            <td style="text-align: right; border: 1px solid #94a3b8; padding: 5px;">{unit_rate:,.2f}</td>
            <td style="text-align: right; border: 1px solid #94a3b8; padding: 5px; font-weight: bold;">{q.base_total_value:,.2f}</td>
        </tr>
        """

    # 🎯 FIX: Injected Full Terms & Conditions into PDF Template
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            @page {{ size: a4 portrait; margin: 15mm; }}
            body {{ font-family: Helvetica, sans-serif; font-size: 9.5pt; color: #1e293b; line-height: 1.4; }}
            .border-table {{ width: 100%; border-collapse: collapse; margin-top: 15px; }}
            .border-table th {{ background-color: #f1f5f9; border: 1px solid #94a3b8; padding: 5px; font-size: 8.5pt; text-align: left; }}
            .tc-section p {{ margin-bottom: 6px; }}
        </style>
    </head>
    <body>
        <table style="width: 100%; border-bottom: 2px solid #2c2a57; padding-bottom: 10px;">
            <tr>
                <td style="font-size: 16pt; font-weight: bold; color: #2c2a57;">AARVI ENCON LIMITED</td>
                <td style="text-align: right; font-size: 9pt;"><b>Ref:</b> AEL/{vendor_name[:6].upper()}-PO/2026-27/{po_number.split('-')[-1]}<br/><b>Date:</b> {date.today().strftime('%d/%m/%Y')}</td>
            </tr>
        </table>
        
        <h2 style="text-align: center; font-size: 12pt; margin-top: 10px; background-color: #f1f5f9; padding: 5px; border: 1px solid #94a3b8;">PURCHASE ORDER</h2>
        
        <table style="width: 100%; margin-top: 10px;">
            <tr>
                <td style="width: 100%; vertical-align: top;">
                    <b>M/s. {vendor_name}</b><br/>
                    {vendor_address}<br/>
                    Cell No.: {vendor_contact}<br/>
                    EMAIL:- {vendor_email}
                </td>
            </tr>
        </table>

        <p style="margin-top: 15px;"><b>Subject: Purchase Order for {winning_quotes[0].product_description if winning_quotes else 'Materials'}.</b></p>
        <p>Dear Sir,<br/>With reference to Quotation Dated {contract_date_str}, and subsequent discussion, we are pleased to inform you that company has decided to place order for the supply of goods with your company.</p>

        <table class="border-table">
            <thead>
                <tr>
                    <th style="width: 8%; text-align: center;">Sr.No.</th>
                    <th style="width: 47%;">Description</th>
                    <th style="width: 15%; text-align: center;">QUANTITY</th>
                    <th style="width: 15%; text-align: right;">RATE UNIT</th>
                    <th style="width: 15%; text-align: right;">Total (Rs.)</th>
                </tr>
            </thead>
            <tbody>
                {table_rows_html}
                <tr>
                    <td colspan="4" style="text-align: right; font-weight: bold; border: 1px solid #94a3b8; padding: 5px;">Basic Total Value</td>
                    <td style="text-align: right; border: 1px solid #94a3b8; padding: 5px;">{base_total:,.2f}</td>
                </tr>
                <tr>
                    <td colspan="4" style="text-align: right; font-weight: bold; border: 1px solid #94a3b8; padding: 5px;">GST Adjustment</td>
                    <td style="text-align: right; border: 1px solid #94a3b8; padding: 5px;">{gst_adj:,.2f}</td>
                </tr>
                <tr>
                    <td colspan="4" style="text-align: right; font-weight: bold; border: 1px solid #94a3b8; padding: 5px; background-color: #f1f5f9;">Net Amount Payable</td>
                    <td style="text-align: right; font-weight: bold; border: 1px solid #94a3b8; padding: 5px; background-color: #f1f5f9;">{net_total:,.2f}</td>
                </tr>
                <tr>
                    <td colspan="5" style="text-align: center; font-style: italic; border: 1px solid #94a3b8; padding: 5px;">({amount_in_words})</td>
                </tr>
            </tbody>
        </table>

        <div class="tc-section" style="margin-top: 15px;">
            <p><b>a) TERMS OF PAYMENTS:</b> {payment_terms}</p>
            <p><b>b) DELIVERY:</b> Time is an essence of this Purchase Order. The material has to be delivered within {time_of_delivery} from the date of issue of PO.</p>
            <p><b>c) PROJECT:</b> {project_name}</p>
            
            <p style="margin-top: 15px;"><b>Our GST Registration no.:</b> 27AAACA3640H1Z0 (Please Confirm the GST No. Before the Preparation of Invoices.)</p>
            <p><b>The placement of order is subject to the following Terms & Conditions:-</b></p>
            
            <p><b>1. PRICE:</b> The cost of Purchase with GST as shown above is Rs. {net_total:,.2f}/- ({amount_in_words}). This is a fixed-price order and no escalation is applicable.</p>
            <p><b>2. QUALITY:</b> If the material supplied is not to the satisfaction of our engineer, then the same has to be replaced without any financial implications.</p>
            <p><b>3. LIQUIDITY DAMAGE: (NOT APPLICABLE)</b> If the supplier fails to deliver all the above-mentioned items within 1 week from the date of PO & Liquidity damages @ 0.5% of the order value per week, subject to a maximum of 5% of the order value will be applicable.</p>
            <p><b>4. TAXES & DUTIES:</b> Prevailing Taxes & Duties (i.e. GST) shall be as shown above.</p>
            <p><b>5. CORRESPONDENCE:</b> All the correspondence pertaining to this order is to be made to:<br/>M/s. Aarvi Encon Ltd.<br/>B-1/603, 6th Floor, Marathon Innova, Marathon Nextgen Complex, G.K. Marg Lower Parel (W), Mumbai -400013. Tel: 022-40499999</p>
            <p><b>6. REFERENCE:</b> Quotation Dated {contract_date_str}.</p>
            <p><b>7. BILLING:</b> Bill to be submitted in 2 sets. Original Bill to be submitted to Head Office Mumbai with copy of Bill to Site for Certification / Verification along with following documents: a) Tax invoice b) Delivery Challan c) P.O. Acceptance Copy.</p>
            <p><b>8. DELIVERY ADDRESS:</b> Contract Person: {site_contact}, Contact No.: {site_phone}. {project_name}. {delivery_address}</p>
            <p><b>9. LEGAL COMPLIANCE:</b> Any disputes or differences arising between the Client and Vendor with respect to this Purchase Order and terms & conditions or any other matter connected with or incidental thereto, it should be exclusive under the arbitration and jurisdiction of the courts of Mumbai.</p>
            
            <p style="margin-top: 15px;">Please acknowledge of the duplicate of this Purchase Order as an acceptance of this Purchase Order.</p>
            <p>Thanking you,</p>
        </div>

        <table style="width: 100%; margin-top: 40px; page-break-inside: avoid;">
            <tr>
                <td style="width: 50%;">
                    Yours faithfully<br/>
                    <b>For M/s. AARVI ENCON LTD.</b><br/><br/><br/><br/>
                    <span style="border-top: 1px solid #000; padding-top: 5px;">AUTHORIZED SIGNATORY</span>
                </td>
                <td style="width: 50%; text-align: right;">
                    Signature & Seal of the Supplier<br/>
                    Accepted & Agreed of the above Said Terms and Conditions<br/><br/><br/><br/>
                    <span style="border-top: 1px solid #000; padding-top: 5px;">ACCEPTED BY SUPPLIER</span>
                </td>
            </tr>
        </table>
    </body>
    </html>
    """

    pdf_buffer = io.BytesIO()
    pisa.CreatePDF(html_content, dest=pdf_buffer)
    pdf_buffer.seek(0)
    return StreamingResponse(
        pdf_buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=Aarvi_PO_{po_number}.pdf"}
    )

# ===================================================================
# 🎯 2. DYNAMIC NATIVE WORD (.DOCX) GENERATOR ENDPOINT
# ===================================================================
@app.get("/api/purchase-orders/{po_number}/download-docx")
def download_word_purchase_order(po_number: str, db: Session = Depends(get_db)):
    po = db.query(models.PurchaseOrder).filter(models.PurchaseOrder.po_number == po_number).first()
    if not po: raise HTTPException(status_code=404, detail="Purchase Order not found.")
    
    ticket = db.query(models.MaterialTicket).filter(models.MaterialTicket.ticket_number == po.ticket_number).first()
    winning_quotes = db.query(models.Quotation).filter(models.Quotation.ticket_number == po.ticket_number, models.Quotation.is_selected == True).all()
    if not winning_quotes:
        winning_quotes = db.query(models.Quotation).filter(models.Quotation.ticket_number == po.ticket_number).all()

    primary_quote = winning_quotes[0] if winning_quotes else None
    
    # 🎯 FIX: Correctly handle 'None' values with strict fallbacks
    vendor_name = getattr(primary_quote, 'vendor_name', None) or "N/A"
    vendor_address = getattr(primary_quote, 'vendor_address', None) or "Address Not Provided"
    vendor_contact = getattr(primary_quote, 'vendor_contact', None) or "N/A"
    vendor_email = getattr(primary_quote, 'vendor_email', None) or "N/A"
    payment_terms = getattr(primary_quote, 'payment_terms', None) or "100% Payment shall be paid after receipt of material at site."
    time_of_delivery = getattr(primary_quote, 'time_of_delivery', None) or "2-3 Days"
    site_contact = getattr(primary_quote, 'site_contact_person', None) or "Site Coordinator"
    site_phone = getattr(primary_quote, 'site_contact_phone', None) or "N/A"
    delivery_address = getattr(primary_quote, 'delivery_address', None) or "Address Pending"
    contract_date = getattr(primary_quote, 'contract_start_date', None)
    contract_date_str = contract_date.strftime('%d/%m/%Y') if contract_date else "Recently Submitted"
    project_name = ticket.project_name if ticket else "N/A"

    base_total = sum(float(q.base_total_value or 0) for q in winning_quotes)
    net_total = sum(float(q.net_amount_payable or q.total_amount or 0) for q in winning_quotes)
    gst_adj = net_total - base_total

    def number_to_words(num):
        if num == 0: return 'Zero'
        ones = ['', 'One ', 'Two ', 'Three ', 'Four ', 'Five ', 'Six ', 'Seven ', 'Eight ', 'Nine ', 'Ten ', 'Eleven ', 'Twelve ', 'Thirteen ', 'Fourteen ', 'Fifteen ', 'Sixteen ', 'Seventeen ', 'Eighteen ', 'Nineteen ']
        tens = ['', '', 'Twenty ', 'Thirty ', 'Forty ', 'Fifty ', 'Sixty ', 'Seventy ', 'Eighty ', 'Ninety ']
        def convert_less_thousand(n):
            s = ''
            if n >= 100: s += ones[int(n // 100)] + 'Hundred '; n %= 100
            if n >= 20: s += tens[int(n // 10)]; n %= 10
            if n > 0: s += ones[int(n)]
            return s
        str_val = ''
        crore = int(num // 10000000); num %= 10000000
        lakh = int(num // 100000); num %= 100000
        thousand = int(num // 1000); num %= 1000
        if crore > 0: str_val += convert_less_thousand(crore) + 'Crore '
        if lakh > 0: str_val += convert_less_thousand(lakh) + 'Lakh '
        if thousand > 0: str_val += convert_less_thousand(thousand) + 'Thousand '
        if num > 0: str_val += convert_less_thousand(num)
        return str_val.strip() + ' Only'

    amount_in_words = f"Rupees {number_to_words(int(round(net_total)))}"

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # 🎯 EMBED LETTERHEAD IMAGE AT THE TOP
    letterhead_path = "assets/letter_head.jpg"
    if os.path.exists(letterhead_path):
        p_lh = doc.add_paragraph()
        p_lh.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_lh.add_run().add_picture(letterhead_path, width=Inches(7.2))

    # Ref & Title
    p_ref = doc.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_ref = p_ref.add_run(f"Ref: AEL/{vendor_name[:6].upper()}-PO/2026-27/{po_number.split('-')[-1]}\nDate: {date.today().strftime('%d/%m/%Y')}")
    run_ref.font.size = Pt(9)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("PURCHASE ORDER")
    run_title.font.bold = True
    run_title.font.size = Pt(12)

    # Vendor Details
    p_vendor = doc.add_paragraph()
    p_vendor.add_run(f"M/s. {vendor_name}\n").bold = True
    p_vendor.add_run(f"{vendor_address}\nCell No.: {vendor_contact} | EMAIL:- {vendor_email}\n").font.size = Pt(9)

    doc.add_paragraph(f"Subject: Purchase Order for {winning_quotes[0].product_description if winning_quotes else 'Materials'}.").runs[0].font.bold = True
    doc.add_paragraph(f"Dear Sir,\nWith reference to Quotation Dated {contract_date_str}, and subsequent discussion, we are pleased to inform you that company has decided to place order for the supply of goods with your company.")

    # Items Table
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    table.autofit = False

    hdr_cells = table.rows[0].cells
    headers = ["Sr.", "Description", "Qty", "Rate", "Total (Rs.)"]
    widths = [Inches(0.5), Inches(3.2), Inches(0.8), Inches(1.0), Inches(1.2)]

    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].width = widths[i]
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    for idx, q in enumerate(winning_quotes, start=1):
        row_cells = table.add_row().cells
        qty = q.quantity or 1
        unit_rate = (q.base_total_value or 0) / qty
        row_cells[0].text = f"0{idx}"
        row_cells[1].text = q.product_description or 'Item'
        row_cells[2].text = f"{qty} Nos"
        row_cells[3].text = f"{unit_rate:,.2f}"
        row_cells[4].text = f"{q.base_total_value:,.2f}"

    r1 = table.add_row().cells
    r1[3].text = "Basic Total Value"
    r1[4].text = f"{base_total:,.2f}"
    
    r2 = table.add_row().cells
    r2[3].text = "GST Adjustment"
    r2[4].text = f"{gst_adj:,.2f}"

    r3 = table.add_row().cells
    r3[3].text = "Net Amount Payable"
    r3[4].text = f"{net_total:,.2f}"
    r3[4].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph(f"({amount_in_words})").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Full Terms
    doc.add_paragraph().add_run(f"a) TERMS OF PAYMENTS: {payment_terms}").font.bold = True
    doc.add_paragraph().add_run(f"b) DELIVERY: Time is an essence of this Purchase Order. The material has to be delivered within {time_of_delivery} from the date of issue of PO.").font.bold = True
    doc.add_paragraph().add_run(f"c) PROJECT: {project_name}").font.bold = True

    doc.add_paragraph("Our GST Registration no.: 27AAACA3640H1Z0 (Please Confirm the GST No. Before the Preparation of Invoices.)")
    doc.add_paragraph("The placement of order is subject to the following Terms & Conditions:-").runs[0].font.bold = True

    terms = [
        f"1. PRICE: The cost of Purchase with GST as shown above is Rs. {net_total:,.2f}/-. This is a fixed-price order and no escalation is applicable.",
        "2. QUALITY: If the material supplied is not to the satisfaction of our engineer, then the same has to be replaced without any financial implications.",
        "3. LIQUIDITY DAMAGE: (NOT APPLICABLE) If the supplier fails to deliver all the above-mentioned items within 1 week from the date of PO & Liquidity damages @ 0.5% of the order value per week, subject to a maximum of 5% of the order value will be applicable.",
        "4. TAXES & DUTIES: Prevailing Taxes & Duties (i.e. GST) shall be as shown above.",
        "5. CORRESPONDENCE: All the correspondence pertaining to this order is to be made to: M/s. Aarvi Encon Ltd., B-1/603, 6th Floor, Marathon Innova, Marathon Nextgen Complex, G.K. Marg Lower Parel (W), Mumbai -400013. Tel: 022-40499999",
        f"6. REFERENCE: Quotation Dated {contract_date_str}.",
        "7. BILLING: Bill to be submitted in 2 sets. Original Bill to be submitted to Head Office Mumbai with copy of Bill to Site for Certification / Verification along with following documents: a) Tax invoice b) Delivery Challan c) P.O. Acceptance Copy.",
        f"8. DELIVERY ADDRESS: Contract Person: {site_contact}, Contact No.: {site_phone}. {project_name}. {delivery_address}",
        "9. LEGAL COMPLIANCE: Any disputes or differences arising between the Client and Vendor with respect to this Purchase Order and terms & conditions or any other matter connected with or incidental thereto, it should be exclusive under the arbitration and jurisdiction of the courts of Mumbai."
    ]
    for term in terms:
        p = doc.add_paragraph(term)
        p.runs[0].font.size = Pt(9)
        
    doc.add_paragraph("Please acknowledge of the duplicate of this Purchase Order as an acceptance of this Purchase Order.\nThanking you,\nYours faithfully")

    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c1, c2 = sig_table.rows[0].cells
    c1.text = "\n\nFor AARVI ENCON LIMITED\n_____________________\nAuthorized Signatory"
    c2.text = "\n\nAccepted By Vendor\n_____________________\nSignature & Seal"

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=Aarvi_PO_{po_number}.docx"}
    )

# -------------------------------------------------------------------
# 🚀 PHASE 2: MASTER PO LEDGER & PROFORMA INVOICE (PI) UPLOAD
# -------------------------------------------------------------------
@app.get("/api/purchase-orders/finalized", response_model=List[dict])
def get_finalized_purchase_orders(db: Session = Depends(get_db)):
    orders = db.query(
        models.PurchaseOrder, 
        models.MaterialTicket
    ).join(
        models.MaterialTicket, models.PurchaseOrder.ticket_number == models.MaterialTicket.ticket_number
    ).filter(models.MaterialTicket.status.in_([
        "Approved", "PI Pending PM Approval", "PI Approved - Sent to Accounts", 
        "Dispatched", "Partially Delivered", "Material Discrepancy Raised", "Delivered - GRN Logged", "Partially Disbursed"
    ])).order_by(models.PurchaseOrder.generated_at.desc()).all()
    
    response = []
    for po_obj, ticket_obj in orders:
        winning_quotes = db.query(models.Quotation).filter(
            models.Quotation.ticket_number == po_obj.ticket_number,
            models.Quotation.is_selected == True
        ).all()
        
        grand_total = sum(q.total_amount for q in winning_quotes)
        primary_quote = winning_quotes[0] if winning_quotes else None
        primary_vendor = primary_quote.vendor_name if primary_quote else "N/A"
        
        items = db.query(models.TicketItem).filter(models.TicketItem.ticket_number == po_obj.ticket_number).all()
        item_list = [
            {
                "desc": i.product_description, 
                "qty": i.quantity, 
                "is_reimbursable": getattr(i, 'is_reimbursable', False)
            } for i in items
        ]
        purposes = list(set([i.purpose for i in items if getattr(i, 'purpose', None)]))
        aggregated_purpose = ", ".join(purposes) if purposes else "General Maintenance"
        
        sm_log = db.query(models.TicketHistory).filter(
            models.TicketHistory.ticket_number == po_obj.ticket_number,
            models.TicketHistory.action_taken == "Explicit Sign-Off Applied",
            models.TicketHistory.remarks.contains("Site Manager")
        ).order_by(models.TicketHistory.timestamp.desc()).first()
        site_manager = sm_log.user_name if sm_log else "Pending / N/A"
        
        pm_log = db.query(models.TicketHistory).filter(
            models.TicketHistory.ticket_number == po_obj.ticket_number,
            models.TicketHistory.action_taken == "Approve"
        ).order_by(models.TicketHistory.timestamp.desc()).first()
        project_manager = pm_log.user_name if pm_log else "Pending / N/A"
        
        response.append({
            "po_number": po_obj.po_number,
            "ticket_number": po_obj.ticket_number,
            "generated_at": po_obj.generated_at.strftime('%d-%b-%Y %I:%M %p') if po_obj.generated_at else "N/A",
            "requisition_date": ticket_obj.created_at.strftime('%d-%b-%Y') if getattr(ticket_obj, 'created_at', None) else "N/A",
            "vendor_name": primary_vendor,
            "vendor_address": getattr(primary_quote, 'vendor_address', 'N/A') if primary_quote else 'N/A',
            "vendor_email": getattr(primary_quote, 'vendor_email', 'N/A') if primary_quote else 'N/A',
            "vendor_contact": getattr(primary_quote, 'vendor_contact', 'N/A') if primary_quote else 'N/A',
            "purpose": aggregated_purpose,
            "grand_total": float(grand_total),
            "project_name": ticket_obj.project_name,
            "project_code": ticket_obj.project_code,
            "site_manager": site_manager,
            "project_manager": project_manager,
            "items": item_list,
            "category": ticket_obj.category,
            "status": ticket_obj.status,
            "po_pdf_url": getattr(po_obj, 'pdf_url', None),
            "signed_po_url": getattr(po_obj, 'signed_po_url', None),
            "invoice_no": getattr(po_obj, 'invoice_no', '') or '',
            "invoice_date": getattr(po_obj, 'invoice_date', '') or '',
            "invoice_remark": getattr(po_obj, 'invoice_remark', '') or '',
            "invoice_duration": getattr(po_obj, 'invoice_duration', '') or '',
            "proforma_invoice_url": getattr(po_obj, 'proforma_invoice_url', None),
            "tax_invoice_no": getattr(po_obj, 'tax_invoice_no', '') or '',
            "tax_invoice_date": getattr(po_obj, 'tax_invoice_date', '') or '',
            "tax_invoice_url": getattr(po_obj, 'tax_invoice_url', None),
            "utr_no": getattr(po_obj, 'utr_no', '') or '',
            "payment_date": getattr(po_obj, 'payment_date', '') or '',
            "payment_remark": getattr(po_obj, 'payment_remark', '') or '',
            "payment_advice_url": getattr(po_obj, 'payment_advice_url', None),
            "disbursed_amount": float(getattr(po_obj, 'disbursed_amount', 0) or 0)
        })
    return response

# 🚀 UPDATED: Uploads the physical PDF attachment directly to Cloudinary
@app.put("/api/purchase-orders/{po_number}/invoice")
async def update_po_invoice_details(
    po_number: str, 
    invoice_no: str = Form(""),
    invoice_date: str = Form(""),
    invoice_remark: str = Form(""),
    invoice_duration: str = Form(""),
    file: Optional[UploadFile] = File(None),
    background_tasks: BackgroundTasks = BackgroundTasks(),
    db: Session = Depends(get_db)
):
    po = db.query(models.PurchaseOrder).filter(models.PurchaseOrder.po_number == po_number).first()
    if not po:
        raise HTTPException(status_code=404, detail="Purchase Order entity not found.")
    
    po.invoice_no = invoice_no
    po.invoice_date = invoice_date
    po.invoice_remark = invoice_remark
    po.invoice_duration = invoice_duration
    
    if file:
        ext = os.path.splitext(file.filename)[1].lower()
        if ext not in [".pdf", ".png", ".jpg", ".jpeg"]:
            raise HTTPException(status_code=400, detail="Only PDF and Image files are allowed.")
            
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"PI_{po_number}_{timestamp}"
        
        try:
            upload_result = cloudinary.uploader.upload(
                file.file, 
                public_id=filename,
                folder="aarvi_invoices",
                resource_type="auto"
            )
            po.proforma_invoice_url = upload_result.get("secure_url")
            
        except Exception as e:
            logger.error(f"Cloudinary Upload Failed: {str(e)}")
            raise HTTPException(status_code=500, detail="Failed to upload document to cloud storage.")
        
        ticket = db.query(models.MaterialTicket).filter(models.MaterialTicket.ticket_number == po.ticket_number).first()
        if ticket:
            ticket.status = "PI Pending PM Approval"
            db.add(models.TicketHistory(
                ticket_number=ticket.ticket_number,
                user_name="Purchase Executive",
                action_taken="Proforma Invoice Uploaded",
                remarks=f"Vendor PI {invoice_no} securely uploaded to cloud and routed to Project Manager for financial clearance."
            ))

            # 🎯 AUTOMATED EMAIL: Alert Project Manager that PI is ready for approval
            pm = db.query(models.User).filter(models.User.id == ticket.assigned_project_manager_id).first()
            if pm and pm.email:
                background_tasks.add_task(
                    send_workflow_email,
                    recipient_email=pm.email,
                    recipient_name=pm.name,
                    subject=f"Action Required: Proforma Invoice Uploaded for {ticket.ticket_number}",
                    ticket_number=ticket.ticket_number,
                    project_name=ticket.project_name,
                    status="PI Pending PM Approval"
                )
    
    db.commit()
    return {"message": "Invoice logging verified, stored in cloud, and saved successfully."}

@app.delete("/api/purchase-orders/{po_number}/invoice-file")
def delete_po_invoice_file(po_number: str, db: Session = Depends(get_db)):
    po = db.query(models.PurchaseOrder).filter(models.PurchaseOrder.po_number == po_number).first()
    if not po:
        raise HTTPException(status_code=404, detail="Purchase Order entity not found.")
    
    if po.proforma_invoice_url:
        try:
            url = po.proforma_invoice_url
            if "/upload/" in url:
                path_after_upload = url.split("/upload/")[1]
                parts = path_after_upload.split('/')
                if parts[0].startswith('v') and parts[0][1:].isdigit():
                    parts = parts[1:]
                
                full_path = "/".join(parts)
                public_id = os.path.splitext(full_path)[0]
                
                cloudinary.uploader.destroy(public_id, resource_type="image")
                cloudinary.uploader.destroy(public_id, resource_type="raw")
                cloudinary.uploader.destroy(full_path, resource_type="raw")
                
        except Exception as e:
            logger.error(f"Failed to delete Cloudinary file: {str(e)}")

        po.proforma_invoice_url = None
        
        ticket = db.query(models.MaterialTicket).filter(models.MaterialTicket.ticket_number == po.ticket_number).first()
        if ticket:
            db.add(models.TicketHistory(
                ticket_number=ticket.ticket_number,
                user_name="Purchase Executive",
                action_taken="Proforma Invoice Attachment Removed",
                remarks=f"Proforma invoice attachment deleted for PO {po_number}."
            ))

    db.commit()
    return {"message": "Attachment deleted successfully from Cloudinary and Database."}    

# -------------------------------------------------------------------
# 🎯 SIGNED PO UPLOAD ENDPOINT
# -------------------------------------------------------------------
@app.put("/api/purchase-orders/{po_number}/signed-po")
async def upload_signed_po_document(
    po_number: str, 
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    po = db.query(models.PurchaseOrder).filter(models.PurchaseOrder.po_number == po_number).first()
    if not po:
        raise HTTPException(status_code=404, detail="Purchase Order not found.")
        
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in [".pdf", ".png", ".jpg", ".jpeg", ".doc", ".docx"]:
        raise HTTPException(status_code=400, detail="Invalid file type.")
        
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"SIGNED_PO_{po_number}_{timestamp}"
    
    try:
        upload_result = cloudinary.uploader.upload(
            file.file, 
            public_id=filename,
            folder="aarvi_signed_pos",
            resource_type="auto"
        )
        po.signed_po_url = upload_result.get("secure_url")
    except Exception as e:
        logger.error(f"Cloudinary Signed PO Upload Failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to upload Signed PO.")
        
    db.commit()
    return {"message": "Signed PO uploaded successfully.", "signed_po_url": po.signed_po_url}

# -------------------------------------------------------------------
# 🧾 TAX INVOICE UPLOAD ENDPOINT
# -------------------------------------------------------------------
@app.put("/api/purchase-orders/{po_number}/tax-invoice")
async def update_po_tax_invoice_details(
    po_number: str, 
    tax_invoice_no: str = Form(""),
    tax_invoice_date: str = Form(""),
    file: Optional[UploadFile] = File(None),
    db: Session = Depends(get_db)
):
    po = db.query(models.PurchaseOrder).filter(models.PurchaseOrder.po_number == po_number).first()
    if not po:
        raise HTTPException(status_code=404, detail="Purchase Order entity not found.")
    
    po.tax_invoice_no = tax_invoice_no
    po.tax_invoice_date = tax_invoice_date
    
    if file:
        ext = os.path.splitext(file.filename)[1].lower()
        if ext not in [".pdf", ".png", ".jpg", ".jpeg"]:
            raise HTTPException(status_code=400, detail="Only PDF and Image files are allowed.")
            
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"TAX_INV_{po_number}_{timestamp}"
        
        try:
            upload_result = cloudinary.uploader.upload(
                file.file, 
                public_id=filename,
                folder="aarvi_tax_invoices",
                resource_type="auto"
            )
            po.tax_invoice_url = upload_result.get("secure_url")
        except Exception as e:
            logger.error(f"Cloudinary Upload Failed: {str(e)}")
            raise HTTPException(status_code=500, detail="Failed to upload Tax Invoice to cloud storage.")
        
        ticket = db.query(models.MaterialTicket).filter(models.MaterialTicket.ticket_number == po.ticket_number).first()
        if ticket:
            db.add(models.TicketHistory(
                ticket_number=ticket.ticket_number,
                user_name="Purchase Executive",
                action_taken="Tax Invoice Uploaded",
                remarks=f"Final Tax Invoice {tax_invoice_no} uploaded to cloud repository."
            ))
    
    db.commit()
    return {"message": "Tax Invoice details saved successfully.", "tax_invoice_url": po.tax_invoice_url}


# --- SYSTEM HEALTH ROUTER ---
@app.get("/")
def connection_ping():
    return {"status": "online", "message": "Aarvi Encon SCM Routing Router fully initialized."}